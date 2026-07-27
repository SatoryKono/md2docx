"""Golden style-critical invariants (no full A.docx required)."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from md2docx.adapters.outbound.docx_writer import DocxWriter
from md2docx.adapters.outbound.markdown_parser import SimpleMarkdownParser
from md2docx.application.convert_md import ConvertMarkdownToDocx
from md2docx.domain.stylespec import PAGE_DEFAULT, RenderOptions

MD = """# ВВЕДЕНИЕ

Основной текст H_2O.

- пункт один
- пункт два

# 1 Раздел

Текст раздела.
"""


def test_golden_margins_and_list_no_numpr(tmp_path: Path):
    out = tmp_path / "g.docx"
    ConvertMarkdownToDocx(SimpleMarkdownParser(), DocxWriter()).execute(MD, out, RenderOptions())
    doc = Document(str(out))
    sec = doc.sections[0]
    assert abs(float(sec.left_margin.mm) - PAGE_DEFAULT["left_mm"]) < 0.6
    assert abs(float(sec.right_margin.mm) - PAGE_DEFAULT["right_mm"]) < 0.6
    assert abs(float(sec.top_margin.mm) - PAGE_DEFAULT["top_mm"]) < 0.6
    assert abs(float(sec.bottom_margin.mm) - PAGE_DEFAULT["bottom_mm"]) < 0.6

    list_paras = [p for p in doc.paragraphs if p.text.startswith("\u2014")]
    assert list_paras
    p0 = list_paras[0]
    assert "\t" in p0.text
    pPr = p0._p.find(qn("w:pPr"))
    if pPr is not None:
        assert pPr.find(qn("w:numPr")) is None

    styles = {s.name for s in doc.styles}
    assert "StructuralHeading" in styles or any("ВВЕДЕНИЕ" in p.text for p in doc.paragraphs)

    found_sub = any(r.font.subscript for p in doc.paragraphs for r in p.runs if r.font.subscript)
    assert found_sub


def test_demo_via_model_cli(tmp_path: Path):
    from md2docx.adapters.inbound.cli import main

    out = tmp_path / "demo.docx"
    assert main(["--demo", "-o", str(out), "-q"]) == 0
    assert out.is_file()
    doc = Document(str(out))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "ВВЕДЕНИЕ" in texts
