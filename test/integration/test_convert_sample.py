from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from md2docx.adapters.outbound.docx_writer import DocxWriter
from md2docx.adapters.outbound.markdown_parser import SimpleMarkdownParser
from md2docx.application.convert_md import ConvertMarkdownToDocx
from md2docx.domain.stylespec import RenderOptions


def test_convert_mini(tmp_path: Path):
    md = """# ВВЕДЕНИЕ

H_2O и E=mc^2.

- первый
- второй

1. шаг
"""
    out = tmp_path / "out.docx"
    ConvertMarkdownToDocx(SimpleMarkdownParser(), DocxWriter()).execute(
        md, out, RenderOptions()
    )
    assert out.is_file()
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("ВВЕДЕНИЕ" in t for t in texts)
    # list: em dash + tab, no numPr
    list_paras = [p for p in doc.paragraphs if p.text.startswith("\u2014")]
    assert list_paras
    p0 = list_paras[0]
    assert "\t" in p0.text
    pPr = p0._p.find(qn("w:pPr"))
    if pPr is not None:
        assert pPr.find(qn("w:numPr")) is None
    # subscript present
    found_sub = False
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.subscript:
                found_sub = True
    assert found_sub
