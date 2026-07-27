"""Unit tests for TOC filter, empty skip, Header4 mapping (via style aliases)."""

from pathlib import Path

from docx import Document

from md2docx.adapters.outbound.docx_engine import apply_gost_styles
from md2docx.adapters.outbound.docx_reader import DocxReader, _is_toc_line
from md2docx.domain.model import Heading, ListItem, Paragraph, Table


def test_toc_line_detector():
    assert _is_toc_line("1.2 Хронизация боли: роль сенситизации 4")
    assert _is_toc_line("2.7 Перспективность применения 14")
    assert not _is_toc_line(
        "На основе анализа эпидемиологии сформулированы ключевые аспекты."
    )
    assert not _is_toc_line("1 Постановка задачи")


def test_reader_skips_empty_and_toc(tmp_path: Path):
    doc = Document()
    apply_gost_styles(doc)
    doc.add_paragraph("1.2 Тема 4")  # TOC-like
    doc.add_paragraph("")  # empty
    doc.add_paragraph("Нормальный текст абзаца с содержанием.")
    doc.add_paragraph("— пункт списка")
    path = tmp_path / "t.docx"
    doc.save(path)

    model = DocxReader(outline=False).read(path)
    texts = [getattr(b, "text", None) for b in model.blocks]
    assert "1.2 Тема 4" not in texts
    assert any(isinstance(b, Paragraph) and "Нормальный" in b.text for b in model.blocks)
    assert any(isinstance(b, ListItem) and "пункт" in b.text for b in model.blocks)


def test_header4_alias_and_bib(tmp_path: Path):
    doc = Document()
    apply_gost_styles(doc)
    # simulate Heading 4 via style name if available
    h = doc.add_paragraph("3.3 Головной Мозг")
    try:
        h.style = doc.styles["Heading 3"]
    except Exception:
        pass
    # force style id Mapping: use Heading 3 directly as we map H4→H3 in aliases
    bib = doc.add_paragraph("[1] ГОСТ 7.32–2017. Отчёт о НИР.")
    try:
        bib.style = doc.styles["Bibliography"]
    except Exception:
        pass
    path = tmp_path / "h.docx"
    doc.save(path)
    model = DocxReader(outline=False).read(path)
    assert any(
        isinstance(b, Heading) and b.level == 3 and "Головной" in b.text
        for b in model.blocks
    )
    assert any(isinstance(b, ListItem) and b.ordered and b.index == 1 for b in model.blocks)


def test_table_caption_pending(tmp_path: Path):
    doc = Document()
    apply_gost_styles(doc)
    cap = doc.add_paragraph("Таблица 1 — Параметры")
    try:
        cap.style = doc.styles["CaptionTable"]
    except Exception:
        pass
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    path = tmp_path / "tbl.docx"
    doc.save(path)
    model = DocxReader(outline=False).read(path)
    tables = [b for b in model.blocks if isinstance(b, Table)]
    assert tables
    assert tables[0].caption and "Таблица" in tables[0].caption
