"""Unit: DocxWriter — все типы блоков, title, multi-section."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from PIL import Image

from md2docx.adapters.outbound.docx_writer import DocxWriter, split_into_sections
from md2docx.domain.model import (
    CodeLine,
    DocumentModel,
    EmptyLine,
    Figure,
    Formula,
    Heading,
    ListItem,
    Paragraph,
    Quote,
    SectionBreak,
    Table,
    TitleMeta,
)
from md2docx.domain.page import PageSetup, page_setup_default
from md2docx.domain.stylespec import RenderOptions


def _png(path: Path, w: int = 200, h: int = 100) -> Path:
    Image.new("RGB", (w, h), (30, 60, 90)).save(path)
    return path


def test_split_into_sections_first_break_sets_setup():
    land = PageSetup(orientation="landscape", width_mm=210, height_mm=297)
    model = DocumentModel(
        blocks=[
            SectionBreak(setup=land),
            Paragraph(text="wide"),
            SectionBreak(setup=page_setup_default()),
            Paragraph(text="port"),
        ],
        default_page=page_setup_default(),
    )
    parts = split_into_sections(model)
    assert len(parts) == 2
    assert parts[0][0].orientation == "landscape"
    assert parts[1][0].orientation == "portrait"
    assert [b.text for b in parts[0][1]] == ["wide"]


def test_split_skips_empty_lines():
    model = DocumentModel(
        blocks=[
            EmptyLine(),
            Paragraph(text="a"),
            EmptyLine(),
            Paragraph(text="b"),
        ]
    )
    parts = split_into_sections(model)
    assert len(parts) == 1
    assert len(parts[0][1]) == 2


def test_write_all_block_types(tmp_path: Path):
    img = _png(tmp_path / "f.png")
    model = DocumentModel(
        title=TitleMeta(
            org="ОРГАНИЗАЦИЯ",
            topic="Тема НИР",
            city_year="Москва – 2026",
        ),
        default_page=page_setup_default(),
        blocks=[
            Heading(level=1, text="ВВЕДЕНИЕ", structural=True),
            Heading(level=1, text="1 Раздел", structural=False),
            Heading(level=2, text="1.1 Подраздел", structural=False),
            Heading(level=3, text="1.1.1 Пункт", structural=False),
            Paragraph(text="Абзац H_2O"),
            ListItem(text="тире", ordered=False),
            ListItem(text="нумерация", ordered=True, index=1),
            Table(
                rows=[["A", "B"], ["1", "2"]],
                caption="Таблица 1 — демо",
            ),
            Figure(caption="Рисунок 1 — схема", path=str(img)),
            Figure(caption="Рисунок без файла", path=None),
            Formula(text="E=mc^2"),
            Quote(text="цитата"),
            CodeLine(text="print(1)"),
            CodeLine(text=""),
            EmptyLine(),
            SectionBreak(
                setup=PageSetup(orientation="landscape", width_mm=210, height_mm=297)
            ),
            Paragraph(text="landscape body"),
            SectionBreak(setup=page_setup_default()),
            Paragraph(text="again portrait"),
        ],
    )
    dest = tmp_path / "all.docx"
    out = DocxWriter().write(model, RenderOptions(page_numbers=True), dest)
    assert out.is_file()
    doc = Document(str(dest))
    assert len(doc.sections) >= 2
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "ВВЕДЕНИЕ" in texts
    assert "Абзац" in texts
    assert "тире" in texts
    assert "цитата" in texts
    assert "print(1)" in texts
    assert "Тема НИР" in texts
    assert len(doc.tables) >= 1
    assert len(doc.inline_shapes) >= 1
    # landscape section present
    orients = []
    for s in doc.sections:
        w, h = float(s.page_width.mm), float(s.page_height.mm)
        orients.append(
            "L"
            if s.orientation == WD_ORIENT.LANDSCAPE or w > h + 0.5
            else "P"
        )
    assert "L" in orients


def test_write_strict_missing_image_raises(tmp_path: Path):
    from md2docx.domain.errors import MediaError

    model = DocumentModel(
        blocks=[
            Figure(caption="Рисунок — x", path=str(tmp_path / "nope.png")),
        ]
    )
    dest = tmp_path / "strict.docx"
    try:
        DocxWriter().write(
            model, RenderOptions(page_numbers=False, strict=True), dest
        )
        raised = False
    except MediaError:
        raised = True
    assert raised


def test_write_table_br_cells(tmp_path: Path):
    model = DocumentModel(
        blocks=[
            Table(rows=[["H"], ["line1<br>line2"]]),
        ]
    )
    dest = tmp_path / "br.docx"
    DocxWriter().write(model, RenderOptions(page_numbers=False), dest)
    doc = Document(str(dest))
    cell_text = doc.tables[0].rows[1].cells[0].text
    assert "line1" in cell_text
    assert "line2" in cell_text
