"""Unit: page helpers в docx_engine."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from PIL import Image

from md2docx.adapters.outbound import docx_engine as eng
from md2docx.domain.page import PageSetup


def test_set_and_read_section_page_roundtrip(tmp_path: Path):
    doc = Document()
    s0 = doc.sections[0]
    eng.set_section_page(
        s0,
        PageSetup(
            orientation="portrait",
            width_mm=210,
            height_mm=297,
            margin_left_mm=25,
            margin_right_mm=15,
            margin_top_mm=18,
            margin_bottom_mm=18,
        ),
    )
    s1 = doc.add_section(WD_SECTION.NEW_PAGE)
    eng.set_section_page(
        s1,
        PageSetup(
            orientation="landscape",
            width_mm=210,
            height_mm=297,
            margin_left_mm=20,
            margin_right_mm=15,
        ),
    )
    path = tmp_path / "p.docx"
    doc.save(str(path))

    doc2 = Document(str(path))
    p0 = eng.read_section_page_setup(doc2.sections[0])
    p1 = eng.read_section_page_setup(doc2.sections[1])
    assert p0.orientation == "portrait"
    assert abs(p0.margin_left_mm - 25) < 0.5
    assert p1.orientation == "landscape"
    assert float(doc2.sections[1].page_width.mm) > float(doc2.sections[1].page_height.mm)


def test_section_text_width_height():
    doc = Document()
    eng.set_section_page(doc.sections[0], PageSetup())
    tw = eng.section_text_width_mm(doc.sections[0])
    th = eng.section_text_height_mm(doc.sections[0])
    assert abs(tw - 165.0) < 1.0  # 210-30-15
    assert abs(th - 257.0) < 2.0  # 297-20-20


def test_add_figure_picture_full_width_and_no_indent(tmp_path: Path):
    img = tmp_path / "i.png"
    Image.new("RGB", (800, 200), (10, 20, 30)).save(img)
    doc = Document()
    eng.apply_gost_styles(doc)
    eng.set_section_page(doc.sections[0], PageSetup())
    shape = eng.add_figure_picture(doc, img, doc.sections[0])
    tw = eng.section_text_width_mm(doc.sections[0])
    assert abs(float(shape.width.mm) - tw) < 1.0
    p = doc.paragraphs[-1]
    fl = p.paragraph_format.first_line_indent
    assert fl is not None and abs(float(fl.mm)) < 0.1


def test_add_figure_picture_tall_capped(tmp_path: Path):
    img = tmp_path / "tall.png"
    Image.new("RGB", (200, 2000), (40, 40, 40)).save(img)
    doc = Document()
    eng.apply_gost_styles(doc)
    eng.set_section_page(doc.sections[0], PageSetup())
    shape = eng.add_figure_picture(doc, img, doc.sections[0])
    th = eng.section_text_height_mm(doc.sections[0])
    assert float(shape.height.mm) <= th + 0.5


def test_restyle_preserves_landscape(tmp_path: Path):
    doc = Document()
    eng.apply_gost_styles(doc)
    eng.set_section_page(doc.sections[0], PageSetup())
    s1 = doc.add_section(WD_SECTION.NEW_PAGE)
    eng.set_section_page(s1, PageSetup(orientation="landscape", width_mm=210, height_mm=297))
    src = tmp_path / "src.docx"
    dst = tmp_path / "dst.docx"
    doc.save(str(src))
    eng.restyle_docx(src, dst, page_numbers=True, prune_styles=True)
    doc2 = Document(str(dst))
    assert len(doc2.sections) >= 2
    s = doc2.sections[1]
    w, h = float(s.page_width.mm), float(s.page_height.mm)
    assert s.orientation == WD_ORIENT.LANDSCAPE or w > h + 0.5


def test_remove_unused_styles_runs(tmp_path: Path):
    doc = Document()
    eng.apply_gost_styles(doc)
    eng.add_paragraph_formatted(doc, "x", style="Normal")
    n = eng.remove_unused_styles(doc)
    assert isinstance(n, int)
    path = tmp_path / "s.docx"
    doc.save(str(path))
    assert path.is_file()
