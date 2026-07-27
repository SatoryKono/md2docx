"""E2E: размер рисунков = полоса набора секции."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from md2docx.adapters.outbound import docx_engine as eng
from md2docx.adapters.outbound.docx_writer import DocxWriter
from md2docx.domain.model import DocumentModel, Figure, Paragraph, SectionBreak, TitleMeta
from md2docx.domain.page import PageSetup, page_setup_default
from md2docx.domain.stylespec import RenderOptions


def _png(path: Path, w: int, h: int, color=(40, 80, 120)) -> Path:
    try:
        from PIL import Image
    except ImportError:
        # minimal 1x1 PNG if Pillow missing
        path.write_bytes(
            bytes.fromhex(
                "89504e470d0a1a0a0000000d4948445200000001000000010802000000907753"
                "de0000000c4944415408d763f8ffff3f0005fe02fe0dc0f1a80000000049454e44ae426082"
            )
        )
        return path
    Image.new("RGB", (w, h), color).save(path)
    return path


def _shape_sizes(doc_path: Path) -> list[tuple[float, float]]:
    doc = Document(str(doc_path))
    out = []
    for sh in doc.inline_shapes:
        out.append((float(sh.width.mm), float(sh.height.mm)))
    return out


def _pic_para_first_line_mm(doc_path: Path) -> list[float | None]:
    """Прямой firstLine indent абзацев с рисунком (None = не задан)."""
    doc = Document(str(doc_path))
    vals: list[float | None] = []
    for p in doc.paragraphs:
        if not p._element.findall(".//" + qn("wp:extent")):
            continue
        pf = p.paragraph_format
        ind = pf.first_line_indent
        vals.append(float(ind.mm) if ind is not None else None)
    return vals


def test_figure_width_equals_text_width(tmp_path: Path):
    img = _png(tmp_path / "w.png", 1200, 400)
    model = DocumentModel(
        title=TitleMeta(),
        default_page=page_setup_default(),  # 210−30−15 = 165
        blocks=[
            Paragraph(text="До"),
            Figure(caption="Рисунок 1 — тест", path=str(img)),
        ],
    )
    dest = tmp_path / "out.docx"
    DocxWriter().write(model, RenderOptions(page_numbers=False), dest)

    sizes = _shape_sizes(dest)
    assert len(sizes) == 1
    w_mm, h_mm = sizes[0]
    doc = Document(str(dest))
    expected = eng.section_text_width_mm(doc.sections[0])
    assert abs(w_mm - expected) < 0.5
    assert abs(expected - 165.0) < 1.0
    # aspect preserved ~ 1200:400 = 3:1
    assert abs(h_mm - w_mm / 3) < 1.0


def test_figure_width_landscape_section(tmp_path: Path):
    img = _png(tmp_path / "w.png", 1600, 400)
    land = PageSetup(
        orientation="landscape",
        width_mm=210,
        height_mm=297,
        margin_left_mm=30,
        margin_right_mm=15,
    )
    model = DocumentModel(
        title=TitleMeta(),
        default_page=page_setup_default(),
        blocks=[
            Paragraph(text="P"),
            SectionBreak(setup=land),
            Figure(caption="Рисунок 1 — landscape", path=str(img)),
        ],
    )
    dest = tmp_path / "land.docx"
    DocxWriter().write(model, RenderOptions(page_numbers=False), dest)

    doc = Document(str(dest))
    assert len(doc.sections) >= 2
    sec = doc.sections[1]
    assert sec.orientation == WD_ORIENT.LANDSCAPE or float(sec.page_width.mm) > float(
        sec.page_height.mm
    )
    expected = eng.section_text_width_mm(sec)  # ~297−30−15 = 252
    w_mm, _ = _shape_sizes(dest)[0]
    assert abs(w_mm - expected) < 0.5
    assert expected > 200  # landscape шире portrait


def test_figure_paragraph_no_first_line_indent(tmp_path: Path):
    """Normal firstLine 1.25 см не должен сдвигать inline-рисунок."""
    img = _png(tmp_path / "w.png", 800, 400)
    model = DocumentModel(
        title=TitleMeta(),
        default_page=page_setup_default(),
        blocks=[Figure(caption="Рисунок 1 — x", path=str(img))],
    )
    dest = tmp_path / "ind.docx"
    DocxWriter().write(model, RenderOptions(page_numbers=False), dest)
    indents = _pic_para_first_line_mm(dest)
    assert indents
    for v in indents:
        assert v is not None
        assert abs(v) < 0.1  # явный 0 mm


def test_tall_figure_capped_to_page_height(tmp_path: Path):
    """Очень высокий рисунок не должен превышать высоту полосы набора."""
    # 1:4 aspect → at 165mm width height would be 660mm >> page
    img = _png(tmp_path / "tall.png", 500, 2000)
    model = DocumentModel(
        title=TitleMeta(),
        default_page=page_setup_default(),
        blocks=[Figure(caption="Рисунок 1 — tall", path=str(img))],
    )
    dest = tmp_path / "tall.docx"
    DocxWriter().write(model, RenderOptions(page_numbers=False), dest)

    doc = Document(str(dest))
    max_h = eng.section_text_height_mm(doc.sections[0])
    w_mm, h_mm = _shape_sizes(dest)[0]
    assert h_mm <= max_h + 0.5
    # scaled down from full text width
    text_w = eng.section_text_width_mm(doc.sections[0])
    assert w_mm < text_w - 1.0
    # aspect ~ 1:4 preserved
    assert abs(h_mm / w_mm - 4.0) < 0.15
