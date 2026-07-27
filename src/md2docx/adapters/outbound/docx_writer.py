"""DocumentWriter adapter: DocumentModel → .docx."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Twips

from md2docx.adapters.outbound import docx_engine as eng
from md2docx.adapters.outbound.docx_list_layout import (
    LIST_LEFT_TWIPS,
)
from md2docx.adapters.outbound.docx_list_layout import (
    WORD_LIST_STYLE_DASH as LIST_STYLE_DASH,
)
from md2docx.adapters.outbound.docx_list_layout import (
    WORD_LIST_STYLE_NUM as LIST_STYLE_NUM,
)
from md2docx.domain.errors import MediaError
from md2docx.domain.list_markers import LIST_MARKER_PREFIX
from md2docx.domain.model import (
    CodeLine,
    DocumentModel,
    EmptyLine,
    Figure,
    Formula,
    Heading,
    ListItem,
    Paragraph,
    Quote,
    SectionBreak,
    Table,
)
from md2docx.domain.page import PageSetup, page_setup_default
from md2docx.domain.stylespec import RenderOptions


def split_into_sections(
    model: DocumentModel,
) -> list[tuple[PageSetup, list]]:
    """Разбить blocks по SectionBreak (break = начало секции с setup)."""
    setup = (model.default_page or page_setup_default()).normalized()
    buf: list = []
    parts: list[tuple[PageSetup, list]] = []

    for b in model.blocks:
        if isinstance(b, SectionBreak):
            if buf:
                parts.append((setup, buf))
                buf = []
            elif not parts:
                # первый break без контента — задаёт setup первой секции
                setup = b.setup.normalized()
                continue
            else:
                parts.append((setup, []))
            setup = b.setup.normalized()
        elif not isinstance(b, EmptyLine):
            buf.append(b)

    parts.append((setup, buf))
    # убрать ведущие пустые секции
    while len(parts) > 1 and not parts[0][1]:
        parts.pop(0)
    if not parts:
        parts = [(page_setup_default(), [])]
    return parts


class DocxWriter:
    def write(
        self,
        model: DocumentModel,
        options: RenderOptions,
        dest: Path,
    ) -> Path:
        self._strict = options.strict
        doc = Document()
        # стили + default page на section 0
        eng.apply_gost_styles(doc, **options.as_style_kwargs())

        parts = split_into_sections(model)
        # section 0 page from model (важнее CLI, если задано)
        eng.set_section_page(doc.sections[0], parts[0][0])
        self._current_section = doc.sections[0]

        self._write_title(doc, model)
        formula_n = 0
        for idx, (setup, content) in enumerate(parts):
            if idx > 0:
                new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
                eng.set_section_page(new_sec, setup)
                self._current_section = new_sec
            else:
                eng.set_section_page(doc.sections[0], setup)
                self._current_section = doc.sections[0]
            for block in content:
                formula_n = self._write_block(doc, block, formula_n)

        if options.page_numbers:
            eng.setup_page_numbers(doc)
        dest = Path(dest)
        doc.save(str(dest))
        return dest

    def _media_problem(self, message: str) -> None:
        print(f"warning: {message}", file=sys.stderr)
        if getattr(self, "_strict", False):
            raise MediaError(message)

    def restyle(
        self,
        source: Path,
        dest: Path,
        options: RenderOptions,
    ) -> Path:
        return Path(
            eng.restyle_docx(
                source,
                dest,
                style_opts=options.as_style_kwargs(),
                page_numbers=options.page_numbers,
            )
        )

    def _write_title(self, doc: Document, model: DocumentModel) -> None:
        t = model.title
        if not (t.org or t.topic or t.city_year):
            return
        if t.org:
            eng.add_paragraph_formatted(doc, t.org, style="TitleOrg")
        eng.add_paragraph_formatted(
            doc,
            "ОТЧЁТ О НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ",
            style="TitleDocType",
        )
        if t.topic:
            eng.add_paragraph_formatted(doc, t.topic, style="TitleTopic")
        if t.city_year:
            eng.add_paragraph_formatted(doc, t.city_year, style="TitleCityYear")

    def _write_block(self, doc: Document, block, formula_n: int) -> int:
        if isinstance(block, SectionBreak):
            return formula_n
        if isinstance(block, Heading):
            if block.structural:
                eng.add_paragraph_formatted(
                    doc, block.text, style="StructuralHeading"
                )
            elif block.level <= 1:
                eng.add_paragraph_formatted(doc, block.text, style="Heading 1")
            elif block.level == 2:
                eng.add_paragraph_formatted(doc, block.text, style="Heading 2")
            else:
                eng.add_paragraph_formatted(doc, block.text, style="Heading 3")
        elif isinstance(block, Paragraph):
            eng.add_paragraph_formatted(
                doc,
                block.text,
                style="Normal",
                spans=getattr(block, "spans", None),
            )
        elif isinstance(block, ListItem):
            p = doc.add_paragraph()
            style_name = LIST_STYLE_NUM if block.ordered else LIST_STYLE_DASH
            try:
                p.style = eng._get_style_by_name(doc, style_name)
            except KeyError:
                p.style = style_name
            if block.ordered:
                n = block.index or 1
                p.add_run(f"{n}) ")
                if getattr(block, "spans", None):
                    eng.add_runs_from_spans(p, block.spans)
                else:
                    eng.add_runs_with_scripts(p, block.text)
            else:
                p.add_run(LIST_MARKER_PREFIX)
                if getattr(block, "spans", None):
                    eng.add_runs_from_spans(p, block.spans)
                else:
                    eng.add_runs_with_scripts(p, block.text)
                try:
                    p.paragraph_format.tab_stops.clear_all()
                except Exception:
                    pass
                p.paragraph_format.tab_stops.add_tab_stop(
                    Twips(LIST_LEFT_TWIPS), alignment=WD_TAB_ALIGNMENT.LEFT
                )
            pPr = p._p.get_or_add_pPr()
            for child in list(pPr):
                if child.tag == qn("w:numPr"):
                    pPr.remove(child)
        elif isinstance(block, Table):
            if block.caption:
                eng.add_paragraph_formatted(
                    doc, block.caption, style="CaptionTable"
                )
            if block.rows:
                cols = max(len(r) for r in block.rows)
                table = doc.add_table(rows=len(block.rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(block.rows):
                    for c_idx in range(cols):
                        cell = table.cell(r_idx, c_idx)
                        text = row[c_idx] if c_idx < len(row) else ""
                        text = text.replace("<br>", "\n").replace("<br/>", "\n")
                        style_name = (
                            "TableHeader" if r_idx == 0 else "TableCell"
                        )
                        eng.fill_table_cell(
                            cell, text, style_name=style_name, doc=doc
                        )
                eng.add_empty_line(doc)
        elif isinstance(block, Figure):
            if block.path:
                img = Path(block.path)
                if not img.is_file():
                    self._media_problem(
                        f"изображение не найдено: {block.path}"
                    )
                else:
                    try:
                        sec = (
                            getattr(self, "_current_section", None)
                            or doc.sections[-1]
                        )
                        eng.add_figure_picture(doc, img, sec)
                    except Exception as exc:
                        self._media_problem(
                            f"не удалось вставить изображение {block.path}: {exc}"
                        )
            cap = eng.add_paragraph_formatted(
                doc, block.caption, style="CaptionFigure"
            )
            if block.path:
                cap.paragraph_format.space_before = Pt(0)
        elif isinstance(block, Formula):
            formula_n += 1
            p = doc.add_paragraph()
            try:
                p.style = eng._get_style_by_name(doc, "Formula")
            except KeyError:
                pass
            eng.add_runs_with_scripts(p, block.text)
            p.add_run(f"\t({formula_n})")
        elif isinstance(block, Quote):
            eng.add_paragraph_formatted(doc, block.text, style="Quote")
        elif isinstance(block, CodeLine):
            p = doc.add_paragraph(block.text if block.text else " ")
            try:
                p.style = eng._get_style_by_name(doc, "CodeBlock")
            except KeyError:
                pass
        elif isinstance(block, EmptyLine):
            eng.add_empty_line(doc)
        return formula_n
