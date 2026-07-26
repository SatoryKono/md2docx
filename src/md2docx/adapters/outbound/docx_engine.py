"""Outbound: python-docx engine (стили ГОСТ + низкоуровневая сборка)."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips

from md2docx.domain.list_markers import (
    LIST_HANGING_TWIPS,
    LIST_LEFT_TWIPS,
    LIST_MARKER_PREFIX,
    LIST_STYLE_DASH,
    LIST_STYLE_NUM,
)
from md2docx.domain.scripts import iter_script_segments, strip_md_inline
from md2docx.domain.structural import is_structural_heading
from md2docx.domain.stylespec import (
    A_FOOTNOTE_PT,
    A_H2_FIRST_LINE_TWIPS,
    A_H3_HANGING_TWIPS,
    A_H3_LEFT_TWIPS,
    A_HEADER1_PT,
    A_HEADER2_PT,
    A_HEADER3_PT,
    A_HEADER4_PT,
    A_LINE_100,
    A_LINE_125,
    A_LINE_150,
    A_NORMAL_FIRST_LINE_TWIPS,
    A_REF_BEFORE_PT,
    A_TABLE_CELL_PAD_PT,
    A_TABLE_HEADER_LEFT_TWIPS,
    A_TOC1_AFTER_PT,
    A_TOC2_AFTER_PT,
    A_TOC2_BEFORE_PT,
    A_TOC2_HANGING_TWIPS,
    A_TOC2_LEFT_TWIPS,
    A_TOC3_AFTER_PT,
    A_TOC3_BEFORE_PT,
    A_TOC3_HANGING_TWIPS,
    A_TOC3_LEFT_TWIPS,
    A_TOC3_PT,
    BODY_PT,
    EMPTY_LINE_PT,
    FIRST_LINE_MM,
    FONT,
    LINE_1_5,
    PAGE_DEFAULT,
    SMALL_PT,
    STYLE_NAMES,
    TABLE_PT,
)

# aliases for remaining code
PAGE = PAGE_DEFAULT
HEADING_COLOR = RGBColor(0x00, 0x00, 0x00)

ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _get_or_add_paragraph_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        pass
    try:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        return doc.styles[name]


def _strip_num_pr(style) -> None:
    """Убрать привязку к Word Numbering (буллеты/автонумерация)."""
    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        return
    for child in list(pPr):
        if child.tag == qn("w:numPr"):
            pPr.remove(child)


def _set_run_font(
    style,
    *,
    font: str,
    size_pt: float,
    bold=False,
    italic=False,
    all_caps=False,
    color: RGBColor | None = HEADING_COLOR,
):
    f = style.font
    f.name = font
    f.size = Pt(size_pt)
    f.bold = bold
    f.italic = italic
    f.all_caps = all_caps
    if color is not None:
        f.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    # Убрать theme-шрифты (иначе Heading тянет Calibri/majorHAnsi)
    for attr in list(rfonts.attrib):
        local = attr.split("}")[-1] if "}" in attr else attr
        if local.endswith("Theme") or "theme" in local.lower():
            del rfonts.attrib[attr]
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)
    # Явно чёрный в XML (сбрасывает themeColor у встроенных Heading)
    if color is not None:
        for child in list(rpr):
            if child.tag == qn("w:color"):
                rpr.remove(child)
        try:
            hex_val = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
        except Exception:
            hex_val = "000000"
        rpr.append(rpr.makeelement(qn("w:color"), {qn("w:val"): hex_val}))


def apply_cell_paragraph_spacing(paragraphs, pad_pt: float = A_TABLE_CELL_PAD_PT) -> None:
    """Отступы в ячейке: 1 абзац → 3pt сверху и снизу;
    несколько → первый 3pt сверху, последний 3pt снизу, остальные 0."""
    n = len(paragraphs)
    if n == 0:
        return
    for i, p in enumerate(paragraphs):
        pf = p.paragraph_format
        if n == 1:
            pf.space_before = Pt(pad_pt)
            pf.space_after = Pt(pad_pt)
        else:
            pf.space_before = Pt(pad_pt if i == 0 else 0)
            pf.space_after = Pt(pad_pt if i == n - 1 else 0)


def fill_table_cell(
    cell,
    text: str,
    *,
    style_name: str = "TableCell",
    pad_pt: float = A_TABLE_CELL_PAD_PT,
) -> None:
    """Заполнить ячейку текстом (\\n = новый абзац) и выставить отступы.
    Поддерживает _подстрочные_ и ^надстрочные индексы."""
    parts = text.split("\n") if text is not None else [""]
    if not parts:
        parts = [""]
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)
    cell.paragraphs[0].text = ""
    cell.paragraphs[0].style = style_name
    add_runs_with_scripts(cell.paragraphs[0], _strip_md_inline(parts[0]))
    for part in parts[1:]:
        p = cell.add_paragraph()
        p.style = style_name
        add_runs_with_scripts(p, _strip_md_inline(part))
    apply_cell_paragraph_spacing(cell.paragraphs, pad_pt=pad_pt)


def add_empty_line(doc: Document) -> None:
    """Пустая строка-отбивка (после таблицы / перед рисунком)."""
    p = doc.add_paragraph("")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = A_LINE_100
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)


def _set_paragraph(
    style,
    *,
    alignment: str,
    first_line_mm: float | None = 0,
    left_mm: float | None = 0,
    right_mm: float | None = 0,
    first_line_twips: int | None = None,
    left_twips: int | None = None,
    hanging_twips: int | None = None,
    before_pt: float = 0,
    after_pt: float = 0,
    line_spacing: float = 1.5,
    keep_with_next: bool = False,
    page_break_before: bool = False,
    outline_level: int | None = None,
):
    pf = style.paragraph_format
    pf.alignment = ALIGN[alignment]
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.keep_with_next = keep_with_next
    pf.page_break_before = page_break_before
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    if hanging_twips is not None:
        left = left_twips if left_twips is not None else hanging_twips
        pf.left_indent = Twips(left)
        pf.first_line_indent = Twips(-hanging_twips)
    elif first_line_twips is not None or left_twips is not None:
        pf.left_indent = Twips(left_twips or 0)
        pf.first_line_indent = Twips(first_line_twips or 0)
    else:
        pf.first_line_indent = Mm(first_line_mm or 0)
        pf.left_indent = Mm(left_mm or 0)

    pf.right_indent = Mm(right_mm or 0)
    if outline_level is not None:
        pPr = style.element.get_or_add_pPr()
        for child in list(pPr):
            if child.tag == qn("w:outlineLvl"):
                pPr.remove(child)
        el = pPr.makeelement(qn("w:outlineLvl"), {qn("w:val"): str(outline_level)})
        pPr.append(el)


def apply_gost_page_setup(doc: Document, page: dict[str, float] | None = None) -> None:
    p = page or PAGE
    for section in doc.sections:
        section.page_width = Mm(p["width_mm"])
        section.page_height = Mm(p["height_mm"])
        section.left_margin = Mm(p["left_mm"])
        section.right_margin = Mm(p["right_mm"])
        section.top_margin = Mm(p["top_mm"])
        section.bottom_margin = Mm(p["bottom_mm"])
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)
        section.different_first_page_header_footer = True


def apply_gost_styles(
    doc: Document,
    *,
    font: str = FONT,
    body_pt: float = BODY_PT,
    table_pt: float = TABLE_PT,
    small_pt: float = SMALL_PT,
    line_spacing: float = LINE_1_5,
    first_line_mm: float = FIRST_LINE_MM,
    page: dict[str, float] | None = None,
) -> Document:
    """Стили по A.docx (заголовки, абзацы, подписи, таблицы, TOC, сноски) + поля ГОСТ."""
    apply_gost_page_setup(doc, page)
    p = page or PAGE
    content_width_mm = p["width_mm"] - p["left_mm"] - p["right_mm"]

    # --- Normal ← A.docx Normal (firstLine 851, both, 1.5, TNR 12) ---
    normal = doc.styles["Normal"]
    _set_run_font(normal, font=font, size_pt=body_pt)
    if abs(first_line_mm - FIRST_LINE_MM) < 0.05:
        _set_paragraph(
            normal,
            alignment="both",
            first_line_twips=A_NORMAL_FIRST_LINE_TWIPS,
            left_twips=0,
            line_spacing=line_spacing,
        )
    else:
        _set_paragraph(
            normal,
            alignment="both",
            first_line_mm=first_line_mm,
            line_spacing=line_spacing,
        )

    # --- Заголовки @Header1…@Header4 (все чёрные) ---
    # @Header1 → StructuralHeading: с новой страницы
    st = _get_or_add_paragraph_style(doc, "StructuralHeading")
    st.base_style = normal
    st.quick_style = True
    _set_run_font(
        st, font=font, size_pt=A_HEADER1_PT, bold=False, all_caps=True, color=HEADING_COLOR
    )
    _set_paragraph(
        st,
        alignment="center",
        first_line_mm=0,
        before_pt=0,
        after_pt=0,
        line_spacing=A_LINE_125,
        keep_with_next=True,
        page_break_before=True,
        outline_level=0,
    )

    # @Header2 → Heading 1: разрыв страницы перед заголовком
    h1 = doc.styles["Heading 1"]
    h1.base_style = normal
    _set_run_font(
        h1, font=font, size_pt=A_HEADER2_PT, bold=False, all_caps=True, color=HEADING_COLOR
    )
    _set_paragraph(
        h1,
        alignment="left",
        first_line_twips=A_H2_FIRST_LINE_TWIPS,
        left_twips=0,
        before_pt=0,
        after_pt=0,
        line_spacing=A_LINE_150,
        keep_with_next=True,
        page_break_before=True,
        outline_level=0,
    )

    # @Header3 → Heading 2: пустая строка перед
    h2 = doc.styles["Heading 2"]
    h2.base_style = normal
    _set_run_font(
        h2, font=font, size_pt=A_HEADER3_PT, bold=False, all_caps=False, color=HEADING_COLOR
    )
    _set_paragraph(
        h2,
        alignment="left",
        left_twips=A_H3_LEFT_TWIPS,
        hanging_twips=A_H3_HANGING_TWIPS,
        before_pt=EMPTY_LINE_PT,
        after_pt=0,
        line_spacing=A_LINE_125,
        keep_with_next=True,
        page_break_before=False,
        outline_level=1,
    )

    # @Header4 → Heading 3: пустая строка перед
    h3 = doc.styles["Heading 3"]
    h3.base_style = normal
    _set_run_font(
        h3,
        font=font,
        size_pt=A_HEADER4_PT,
        bold=True,
        italic=True,
        all_caps=False,
        color=HEADING_COLOR,
    )
    _set_paragraph(
        h3,
        alignment="left",
        left_twips=A_H3_LEFT_TWIPS,
        hanging_twips=A_H3_HANGING_TWIPS,
        before_pt=EMPTY_LINE_PT,
        after_pt=0,
        line_spacing=A_LINE_125,
        keep_with_next=True,
        page_break_before=False,
        outline_level=2,
    )

    # --- Подписи: @Title.Table / @Title.Picture ---
    # пустая строка ПЕРЕД названием таблицы; ПОСЛЕ таблицы — add_empty_line() в контенте;
    # пустая строка ПЕРЕД рисунком (caption/image) + ПОСЛЕ подписи рисунка
    cap_t = _get_or_add_paragraph_style(doc, "CaptionTable")
    cap_t.base_style = normal
    _set_run_font(cap_t, font=font, size_pt=body_pt, color=HEADING_COLOR)
    _set_paragraph(
        cap_t,
        alignment="left",
        first_line_mm=0,
        before_pt=EMPTY_LINE_PT,
        after_pt=0,
        line_spacing=A_LINE_125,
        keep_with_next=True,
    )

    cap_f = _get_or_add_paragraph_style(doc, "CaptionFigure")
    cap_f.base_style = normal
    _set_run_font(cap_f, font=font, size_pt=body_pt, color=HEADING_COLOR)
    _set_paragraph(
        cap_f,
        alignment="center",
        first_line_mm=0,
        before_pt=EMPTY_LINE_PT,
        after_pt=EMPTY_LINE_PT,
        line_spacing=A_LINE_125,
        keep_with_next=False,
    )

    # --- Таблицы: @Table.Text / @Table.Header-C / @Table.Header-L ---
    # default 3pt/3pt для одного абзаца; multi-para — fill_table_cell / apply_cell_paragraph_spacing
    tc = _get_or_add_paragraph_style(doc, "TableCell")
    tc.base_style = normal
    _set_run_font(tc, font=font, size_pt=table_pt, color=HEADING_COLOR)
    _set_paragraph(
        tc,
        alignment="center",
        first_line_mm=0,
        before_pt=A_TABLE_CELL_PAD_PT,
        after_pt=A_TABLE_CELL_PAD_PT,
        line_spacing=A_LINE_100,
    )

    th = _get_or_add_paragraph_style(doc, "TableHeader")
    th.base_style = normal
    _set_run_font(th, font=font, size_pt=table_pt, bold=False, color=HEADING_COLOR)
    _set_paragraph(
        th,
        alignment="center",
        first_line_mm=0,
        before_pt=A_TABLE_CELL_PAD_PT,
        after_pt=A_TABLE_CELL_PAD_PT,
        line_spacing=A_LINE_100,
        keep_with_next=True,
    )

    thl = _get_or_add_paragraph_style(doc, "TableHeaderLeft")
    thl.base_style = normal
    _set_run_font(thl, font=font, size_pt=table_pt, bold=False, color=HEADING_COLOR)
    _set_paragraph(
        thl,
        alignment="left",
        first_line_twips=0,
        left_twips=A_TABLE_HEADER_LEFT_TWIPS,
        before_pt=A_TABLE_CELL_PAD_PT,
        after_pt=A_TABLE_CELL_PAD_PT,
        line_spacing=A_LINE_100,
        keep_with_next=True,
    )

    # --- Формула (в A.docx отдельного стиля нет — center + правый таб номера) ---
    formula = _get_or_add_paragraph_style(doc, "Formula")
    formula.base_style = normal
    _set_run_font(formula, font=font, size_pt=body_pt)
    _set_paragraph(
        formula,
        alignment="center",
        first_line_mm=0,
        before_pt=6,
        after_pt=6,
        line_spacing=line_spacing,
    )
    try:
        formula.paragraph_format.tab_stops.clear_all()
    except Exception:
        pass
    formula.paragraph_format.tab_stops.add_tab_stop(
        Mm(content_width_mm), alignment=WD_TAB_ALIGNMENT.RIGHT
    )

    # --- Списки (маркер/номер в тексте, БЕЗ w:numPr — иначе Word рисует •) ---
    # GostListDash: ненумерованный, префикс «—\t»
    ls_dash = _get_or_add_paragraph_style(doc, LIST_STYLE_DASH)
    ls_dash.base_style = normal
    try:
        ls_dash.name = "ГОСТ список (тире)"
    except Exception:
        pass
    _set_run_font(ls_dash, font=font, size_pt=body_pt, color=HEADING_COLOR)
    _set_paragraph(
        ls_dash,
        alignment="both",
        left_twips=LIST_LEFT_TWIPS,
        hanging_twips=LIST_HANGING_TWIPS,
        line_spacing=line_spacing,
    )
    _strip_num_pr(ls_dash)
    # табуляция после «—» → текст с позиции left indent
    try:
        ls_dash.paragraph_format.tab_stops.clear_all()
    except Exception:
        pass
    ls_dash.paragraph_format.tab_stops.add_tab_stop(
        Twips(LIST_LEFT_TWIPS), alignment=WD_TAB_ALIGNMENT.LEFT
    )

    # GostListNumber: нумерованный, префикс «1) » в тексте
    ls_num = _get_or_add_paragraph_style(doc, LIST_STYLE_NUM)
    ls_num.base_style = normal
    try:
        ls_num.name = "ГОСТ список (нумерация)"
    except Exception:
        pass
    _set_run_font(ls_num, font=font, size_pt=body_pt, color=HEADING_COLOR)
    _set_paragraph(
        ls_num,
        alignment="both",
        left_twips=LIST_LEFT_TWIPS,
        hanging_twips=LIST_HANGING_TWIPS,
        line_spacing=line_spacing,
    )
    _strip_num_pr(ls_num)

    # На всякий случай снять буллеты со встроенных List Bullet / List Number
    for builtin in ("List Bullet", "List Number", "List Paragraph"):
        try:
            _strip_num_pr(doc.styles[builtin])
        except KeyError:
            pass

    # --- Библиография ← @Paragraph s1.0_i0.0 (ссылки [1] …) ---
    bib = _get_or_add_paragraph_style(doc, "Bibliography")
    bib.base_style = normal
    _set_run_font(bib, font=font, size_pt=body_pt)
    _set_paragraph(
        bib,
        alignment="both",
        first_line_mm=0,
        before_pt=A_REF_BEFORE_PT,
        after_pt=A_REF_BEFORE_PT,
        line_spacing=A_LINE_100,
    )

    # --- Цитата ← Quote (курсив, как Normal) ---
    quote = _get_or_add_paragraph_style(doc, "Quote")
    quote.base_style = normal
    _set_run_font(quote, font=font, size_pt=body_pt, italic=True)
    if abs(first_line_mm - FIRST_LINE_MM) < 0.05:
        _set_paragraph(
            quote,
            alignment="both",
            first_line_twips=A_NORMAL_FIRST_LINE_TWIPS,
            line_spacing=line_spacing,
        )
    else:
        _set_paragraph(
            quote,
            alignment="both",
            first_line_mm=first_line_mm,
            line_spacing=line_spacing,
        )

    # --- Сноска ← @Foot-note (10 pt, без абзацного отступа) ---
    fn = _get_or_add_paragraph_style(doc, "FootnoteText")
    fn.base_style = normal
    _set_run_font(fn, font=font, size_pt=A_FOOTNOTE_PT)
    _set_paragraph(
        fn,
        alignment="left",
        first_line_mm=0,
        before_pt=0,
        after_pt=0,
        line_spacing=A_LINE_100,
    )
    # встроенный стиль Word «footnote text», если есть
    try:
        fnt = doc.styles["footnote text"]
        _set_run_font(fnt, font=font, size_pt=A_FOOTNOTE_PT)
        _set_paragraph(
            fnt,
            alignment="left",
            first_line_mm=0,
            line_spacing=A_LINE_100,
        )
    except KeyError:
        pass

    # --- Оглавление toc 1…3 (как в A.docx) ---
    toc1 = _get_or_add_paragraph_style(doc, "toc 1")
    toc1.base_style = normal
    _set_run_font(toc1, font=font, size_pt=body_pt)
    _set_paragraph(
        toc1,
        alignment="left",
        first_line_mm=0,
        before_pt=0,
        after_pt=A_TOC1_AFTER_PT,
        line_spacing=line_spacing,
    )

    toc2 = _get_or_add_paragraph_style(doc, "toc 2")
    toc2.base_style = normal
    _set_run_font(toc2, font=font, size_pt=body_pt, all_caps=True)
    _set_paragraph(
        toc2,
        alignment="left",
        left_twips=A_TOC2_LEFT_TWIPS,
        hanging_twips=A_TOC2_HANGING_TWIPS,
        before_pt=A_TOC2_BEFORE_PT,
        after_pt=A_TOC2_AFTER_PT,
        line_spacing=line_spacing,
    )

    toc3 = _get_or_add_paragraph_style(doc, "toc 3")
    toc3.base_style = normal
    _set_run_font(toc3, font=font, size_pt=A_TOC3_PT)
    _set_paragraph(
        toc3,
        alignment="left",
        left_twips=A_TOC3_LEFT_TWIPS,
        hanging_twips=A_TOC3_HANGING_TWIPS,
        before_pt=A_TOC3_BEFORE_PT,
        after_pt=A_TOC3_AFTER_PT,
        line_spacing=A_LINE_125,
    )

    # --- Колонтитул / номер страницы ---
    foot = _get_or_add_paragraph_style(doc, "FooterPageNumber")
    foot.base_style = normal
    _set_run_font(foot, font=font, size_pt=body_pt)
    _set_paragraph(
        foot,
        alignment="center",
        first_line_mm=0,
        line_spacing=A_LINE_100,
    )
    try:
        footer_style = doc.styles["Footer"]
        _set_run_font(footer_style, font=font, size_pt=body_pt)
        _set_paragraph(
            footer_style,
            alignment="center",
            first_line_mm=0,
            line_spacing=A_LINE_100,
        )
    except KeyError:
        pass

    # --- Титул (A.docx @Frontpage.*) + служебные ---
    title_specs = [
        ("TitleOrg", body_pt, True, False, "center", 0, 0, line_spacing),
        ("TitleDocType", body_pt, True, True, "center", 24, 12, line_spacing),
        ("TitleTopic", body_pt, False, True, "center", 12, 12, line_spacing),
        ("TitleMeta", body_pt, False, False, "left", 0, 0, line_spacing),
        ("TitleCityYear", body_pt, False, False, "center", 0, 0, line_spacing),
        ("CodeBlock", small_pt, False, False, "left", 6, 6, A_LINE_100),
    ]
    for style_id, size, caps, bold, align, before, after, ls in title_specs:
        est = _get_or_add_paragraph_style(doc, style_id)
        est.base_style = normal
        use_font = "Courier New" if style_id == "CodeBlock" else font
        _set_run_font(
            est, font=use_font, size_pt=size, bold=bold, all_caps=caps
        )
        _set_paragraph(
            est,
            alignment=align,
            first_line_mm=0,
            before_pt=before,
            after_pt=after,
            line_spacing=ls,
        )

    return doc


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def setup_page_numbers(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    try:
        fp.style = doc.styles["FooterPageNumber"]
    except KeyError:
        pass
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # очистить старый текст
    fp.clear()
    add_page_number_field(fp)


def _strip_md_inline(text: str) -> str:
    return strip_md_inline(text)


def add_runs_with_scripts(paragraph, text: str) -> None:
    """Добавить runs с поддержкой подстрочных (_) и надстрочных (^) индексов."""
    if not text:
        return
    for kind, content in iter_script_segments(text):
        if not content and kind != "plain":
            continue
        if kind == "plain":
            if content:
                paragraph.add_run(content)
        elif kind == "sub":
            run = paragraph.add_run(content)
            run.font.subscript = True
        elif kind == "super":
            run = paragraph.add_run(content)
            run.font.superscript = True


def add_paragraph_formatted(doc: Document, text: str, *, style: str | None = None) -> Any:
    clean = strip_md_inline(text) if text else ""
    kwargs = {}
    if style:
        kwargs["style"] = style
    p = doc.add_paragraph(**kwargs)
    add_runs_with_scripts(p, clean)
    return p


def restyle_docx(
    input_path: str | Path,
    output_path: str | Path,
    *,
    style_opts: dict[str, Any] | None = None,
    page_numbers: bool = True,
) -> str:
    """Открыть существующий .docx, наложить стили/поля ГОСТ, сохранить."""
    doc = Document(str(input_path))
    apply_gost_styles(doc, **(style_opts or {}))
    if page_numbers:
        setup_page_numbers(doc)
    out = str(output_path)
    doc.save(out)
    return out


def build_demo(
    path: str = "gost-demo.docx",
    *,
    style_opts: dict[str, Any] | None = None,
    page_numbers: bool = True,
) -> str:
    style_opts = style_opts or {}
    doc = Document()
    apply_gost_styles(doc, **style_opts)

    doc.add_paragraph("МИНИСТЕРСТВО … / ОРГАНИЗАЦИЯ", style="TitleOrg")
    doc.add_paragraph(
        "ОТЧЁТ О НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ", style="TitleDocType"
    )
    doc.add_paragraph("Наименование темы НИР", style="TitleTopic")
    doc.add_paragraph("Город – 2026", style="TitleCityYear")

    doc.add_paragraph("ВВЕДЕНИЕ", style="StructuralHeading")
    body_pt = style_opts.get("body_pt", BODY_PT)
    doc.add_paragraph(
        "Текст введения оформляется стилем «Основной текст»: выравнивание по ширине, "
        f"абзацный отступ 1,25 см, интервал 1,5, Times New Roman {body_pt:g} pt."
    )

    doc.add_paragraph("1 Постановка задачи", style="Heading 1")
    doc.add_paragraph("1.1 Исходные данные", style="Heading 2")
    doc.add_paragraph(
        "Пример абзаца основного текста отчёта о научно-исследовательской работе."
    )

    doc.add_paragraph("Таблица 1 — Пример названия таблицы", style="CaptionTable")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    fill_table_cell(table.cell(0, 0), "Параметр", style_name="TableHeader")
    fill_table_cell(table.cell(0, 1), "Значение", style_name="TableHeader")
    fill_table_cell(table.cell(1, 0), "Поле левое", style_name="TableCell")
    fill_table_cell(
        table.cell(1, 1),
        "30 мм\n(минимум по ГОСТ)",
        style_name="TableCell",
    )
    add_empty_line(doc)  # после таблицы

    doc.add_paragraph("Рисунок 1 — Пример подписи", style="CaptionFigure")

    p = doc.add_paragraph(style="Formula")
    p.add_run("E = mc")
    p.add_run("2").font.superscript = True
    p.add_run("\t(1)")

    doc.add_paragraph("ЗАКЛЮЧЕНИЕ", style="StructuralHeading")
    doc.add_paragraph("Краткое изложение результатов работы.")

    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", style="StructuralHeading")
    doc.add_paragraph(
        "1. ГОСТ 7.32–2017. Отчёт о научно-исследовательской работе. "
        "Структура и правила оформления. — М., 2017.",
        style="Bibliography",
    )

    if page_numbers:
        setup_page_numbers(doc)

    doc.save(path)
    return path


